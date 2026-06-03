from pathlib import Path
import zipfile

from docx import Document
from docx.parts.hdrftr import FooterPart, HeaderPart

from xlent_scanner.patch import patch_docx, strip_docx_annotations, strip_pptx_annotations, strip_xlsx_annotations


def test_patch_docx_without_header_footer_does_not_create_defaults(
    tmp_path: Path, monkeypatch
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Kunde Ola Nordmann har konto 12345678901.")
    doc.save(source)

    def fail_default_xml(cls):
        raise AssertionError("default header/footer template should not be loaded")

    monkeypatch.setattr(HeaderPart, "_default_header_xml", classmethod(fail_default_xml))
    monkeypatch.setattr(FooterPart, "_default_footer_xml", classmethod(fail_default_xml))

    patch_docx(source, {"12345678901": "[ANONYMISERT]"}, output)

    patched = Document(output)
    text = "\n".join(p.text for p in patched.paragraphs)
    assert "12345678901" not in text
    assert "[ANONYMISERT]" in text


def test_patch_docx_deep_scan_numeric_replacements_stay_xml_safe(tmp_path: Path) -> None:
    source = tmp_path / "budget.docx"
    output = tmp_path / "budget_out.docx"

    doc = Document()
    table = doc.add_table(rows=4, cols=4)
    rows = [
        ("Item", "Cost (NOK)", "Amount", "Total Cost (NOK)"),
        ("Bread", "30", "2", "60"),
        ("Milk", "20", "1", "20"),
        ("Eggs", "10", "10", "100"),
    ]
    for row, values in zip(table.rows, rows, strict=True):
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
    doc.save(source)

    patch_docx(
        source,
        {
            "30": "[ANONYMISERT]",
            "20": "[ANONYMISERT]",
            "10": "[ANONYMISERT]",
            "100": "[ANONYMISERT]",
            "0": "<PNR 1>",
            "1": "<Person A>",
            "2": "<Tlf 3>",
        },
        output,
    )

    patched = Document(output)
    text = "\n".join(
        cell.text for table in patched.tables for row in table.rows for cell in row.cells
    )
    assert "\x00" not in text
    assert "29<" not in text
    assert "30<" not in text
    assert "20<" not in text
    assert "10<" not in text
    assert "100<" not in text
    assert "60" in text


def test_strip_docx_annotations_removes_comment_parts_and_refs(tmp_path: Path) -> None:
    docx = tmp_path / "commented.docx"
    with zipfile.ZipFile(docx, "w") as z:
        z.writestr(
            "[Content_Types].xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>
</Types>""",
        )
        z.writestr(
            "word/document.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body><w:p><w:commentRangeStart w:id="0"/><w:r><w:t>Text</w:t></w:r><w:commentRangeEnd w:id="0"/><w:r><w:commentReference w:id="0"/></w:r></w:p></w:body>
</w:document>""",
        )
        z.writestr(
            "word/_rels/document.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>
</Relationships>""",
        )
        z.writestr("word/comments.xml", "<comments/>")

    strip_docx_annotations(docx)

    with zipfile.ZipFile(docx) as z:
        names = set(z.namelist())
        document_xml = z.read("word/document.xml").decode("utf-8")
        rels_xml = z.read("word/_rels/document.xml.rels").decode("utf-8")
    assert "word/comments.xml" not in names
    assert "commentRangeStart" not in document_xml
    assert "commentReference" not in document_xml
    assert "comments" not in rels_xml


def test_strip_pptx_annotations_removes_notes_and_comments(tmp_path: Path) -> None:
    pptx = tmp_path / "notes.pptx"
    with zipfile.ZipFile(pptx, "w") as z:
        z.writestr(
            "[Content_Types].xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Override PartName="/ppt/notesSlides/notesSlide1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.notesSlide+xml"/>
  <Override PartName="/ppt/comments/comment1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.comments+xml"/>
</Types>""",
        )
        z.writestr("ppt/slides/slide1.xml", "<p:sld xmlns:p=\"http://schemas.openxmlformats.org/presentationml/2006/main\"/>")
        z.writestr(
            "ppt/slides/_rels/slide1.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/notesSlide" Target="../notesSlides/notesSlide1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="../comments/comment1.xml"/>
</Relationships>""",
        )
        z.writestr("ppt/notesSlides/notesSlide1.xml", "<notes/>")
        z.writestr("ppt/comments/comment1.xml", "<comments/>")

    strip_pptx_annotations(pptx)

    with zipfile.ZipFile(pptx) as z:
        names = set(z.namelist())
        rels_xml = z.read("ppt/slides/_rels/slide1.xml.rels").decode("utf-8")
    assert "ppt/notesSlides/notesSlide1.xml" not in names
    assert "ppt/comments/comment1.xml" not in names
    assert "notesSlide" not in rels_xml
    assert "comments" not in rels_xml


def test_strip_xlsx_annotations_removes_comments(tmp_path: Path) -> None:
    xlsx = tmp_path / "comments.xlsx"
    with zipfile.ZipFile(xlsx, "w") as z:
        z.writestr(
            "[Content_Types].xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Override PartName="/xl/comments1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.comments+xml"/>
</Types>""",
        )
        z.writestr(
            "xl/worksheets/sheet1.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"><legacyDrawing r:id="rId1" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/></worksheet>""",
        )
        z.writestr(
            "xl/worksheets/_rels/sheet1.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="../comments1.xml"/>
</Relationships>""",
        )
        z.writestr("xl/comments1.xml", "<comments/>")

    strip_xlsx_annotations(xlsx)

    with zipfile.ZipFile(xlsx) as z:
        names = set(z.namelist())
        sheet_xml = z.read("xl/worksheets/sheet1.xml").decode("utf-8")
        rels_xml = z.read("xl/worksheets/_rels/sheet1.xml.rels").decode("utf-8")
    assert "xl/comments1.xml" not in names
    assert "legacyDrawing" not in sheet_xml
    assert "comments" not in rels_xml
