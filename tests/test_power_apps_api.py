from __future__ import annotations

import base64
from pathlib import Path

from docx import Document

from xlent_scanner import app as app_module
from xlent_scanner.models import Finding, ScanResult, SuppressedFinding


def _fake_result() -> ScanResult:
    return ScanResult(
        file_name="Power Apps tekst",
        file_size=0,
        text_length=42,
        text_preview="PREVIEW_SECRET_NOT_RETURNED",
        findings=[
            Finding(
                category="e-post",
                text="masked@example.com",
                context="kort kontekst",
                severity="gul",
                raw_text="RAW_SECRET_NOT_RETURNED",
            )
        ],
        risk_level="gul",
        risk_summary="Funn oppdaget",
        recommended_action="Kontroller funnene.",
        original_text="FULL_SECRET_NOT_RETURNED",
        language="nb",
    )


def test_api_scan_text_uses_separate_state_and_omits_original_text(monkeypatch):
    monkeypatch.delenv("XLENT_SCANNER_API_KEY", raising=False)
    monkeypatch.setattr(app_module, "scan_text", lambda *args, **kwargs: _fake_result())
    app_module.app_state.api_scan_results.clear()
    sentinel = object()
    app_module.app_state.last_result = sentinel

    client = app_module.flask_app.test_client()
    response = client.post("/api/scan-text", json={"text": "test", "language": "nb"})

    assert response.status_code == 200
    data = response.get_json()
    assert data["ok"] is True
    assert data["scan_id"]
    assert "original_text" not in data
    assert "text_preview" not in data
    assert "FULL_SECRET_NOT_RETURNED" not in str(data)
    assert "RAW_SECRET_NOT_RETURNED" not in str(data)
    assert app_module.app_state.last_result is sentinel


def test_api_scan_text_accepts_scan_profile_and_optional_suppressed(monkeypatch):
    monkeypatch.delenv("XLENT_SCANNER_API_KEY", raising=False)
    captured = {}

    def fake_scan_text(*args, **kwargs):
        captured.update(kwargs)
        result = _fake_result()
        result.suppressed_findings = [
            SuppressedFinding(
                category="telefonnummer",
                text="pp. 4662-4666",
                reason="bibliografisk DOI/ISBN/ISSN/sidekontekst",
            )
        ]
        return result

    monkeypatch.setattr(app_module, "scan_text", fake_scan_text)
    app_module.app_state.api_scan_results.clear()

    client = app_module.flask_app.test_client()
    response = client.post(
        "/api/scan-text",
        json={
            "text": "test",
            "language": "en",
            "scan_profile": "technical",
            "include_suppressed": True,
        },
    )

    assert response.status_code == 200
    assert captured["scan_profile"] == "technical"
    data = response.get_json()
    assert data["suppressed_findings"][0]["text"] == "pp. 4662-4666"


def test_api_scan_text_requires_key_when_configured(monkeypatch):
    monkeypatch.setenv("XLENT_SCANNER_API_KEY", "secret-key")
    monkeypatch.setattr(app_module, "scan_text", lambda *args, **kwargs: _fake_result())
    app_module.app_state.api_scan_results.clear()
    client = app_module.flask_app.test_client()

    missing_key = client.post("/api/scan-text", json={"text": "test"})
    assert missing_key.status_code == 401

    with_key = client.post(
        "/api/scan-text",
        json={"text": "test"},
        headers={"X-API-Key": "secret-key"},
    )
    assert with_key.status_code == 200
    assert with_key.get_json()["ok"] is True


def test_api_scan_file_accepts_base64_and_omits_original_text(monkeypatch):
    monkeypatch.delenv("XLENT_SCANNER_API_KEY", raising=False)
    monkeypatch.setattr(app_module, "scan_file", lambda *args, **kwargs: _fake_result())
    app_module.app_state.api_scan_results.clear()
    client = app_module.flask_app.test_client()

    response = client.post(
        "/api/scan-file",
        json={
            "file_name": "kunde.txt",
            "content_base64": base64.b64encode(b"hello").decode("ascii"),
            "language": "nb",
        },
    )

    assert response.status_code == 200
    data = response.get_json()
    assert data["file_name"] == "kunde.txt"
    assert "original_text" not in data
    assert "FULL_SECRET_NOT_RETURNED" not in str(data)


def test_api_scan_file_rejects_invalid_base64(monkeypatch):
    monkeypatch.delenv("XLENT_SCANNER_API_KEY", raising=False)
    client = app_module.flask_app.test_client()

    response = client.post(
        "/api/scan-file",
        json={"file_name": "kunde.txt", "content_base64": "not base64!!!"},
    )

    assert response.status_code == 400
    assert response.get_json()["error_code"] == "invalid_base64"


def test_api_refuses_network_bind_without_api_key(monkeypatch):
    monkeypatch.delenv("XLENT_SCANNER_API_KEY", raising=False)

    try:
        app_module._validate_api_bind("0.0.0.0")
    except RuntimeError as exc:
        assert "XLENT_SCANNER_API_KEY" in str(exc)
    else:
        raise AssertionError("Expected network bind without API key to fail")

    app_module._validate_api_bind("127.0.0.1")

    monkeypatch.setenv("XLENT_SCANNER_API_KEY", "secret-key")
    app_module._validate_api_bind("0.0.0.0")


def test_settings_export_import_excludes_scan_history_and_document_text(monkeypatch, tmp_path):
    monkeypatch.setenv("APPDATA", str(tmp_path))
    client = app_module.flask_app.test_client()

    exported = client.post(
        "/settings/export",
        json={"browser_settings": {"theme": "light", "language": "nb"}},
    )

    assert exported.status_code == 200
    payload = exported.get_json()
    assert payload["ok"] is True
    assert payload["format"] == "xlent-scanner-settings"
    assert payload["blacklist"] == []
    assert "custom_patterns_toml" in payload
    assert "scan_history" not in payload
    assert "original_text" not in payload

    profile = {
        "format": "xlent-scanner-settings",
        "browser_settings": {"theme": "dark", "language": "en"},
        "whitelist": ["safe@example.com"],
        "blacklist": ["Project Raven"],
        "ignore_toml": 'email_domains = ["xlent.no"]\nnames = ["Test User"]\n',
        "custom_patterns_toml": """\
[[patterns]]
name = "Sak"
regex = 'SAK-\\d{3}'
severity = "gul"
""",
    }
    imported = client.post("/settings/import", json=profile)

    assert imported.status_code == 200
    data = imported.get_json()
    assert data["ok"] is True
    assert data["browser_settings"]["theme"] == "dark"
    assert data["whitelist"] == ["safe@example.com"]
    assert data["blacklist"] == ["Project Raven"]
    assert "Test User" in data["ignore_toml"]
    assert "SAK-" in data["custom_patterns_toml"]


def test_gui_deep_scan_status_route_accepts_specific_job_id(monkeypatch):
    from xlent_scanner import deep_scanner

    monkeypatch.setattr(
        deep_scanner,
        "get_deep_scan_status",
        lambda job_id=None: {
            "job_id": job_id,
            "status": "done",
            "progress": "Ferdig",
            "findings": [],
        } if job_id == "abc123" else {},
    )

    client = app_module.flask_app.test_client()

    found = client.get("/ollama/deep-scan/status/abc123")
    missing = client.get("/ollama/deep-scan/status/missing")

    assert found.status_code == 200
    assert found.get_json()["job_id"] == "abc123"
    assert missing.status_code == 404


def test_openapi_json_and_swagger_docs_are_available():
    client = app_module.flask_app.test_client()

    spec_response = client.get("/api/openapi.json")
    docs_response = client.get("/api/docs")

    assert spec_response.status_code == 200
    spec = spec_response.get_json()
    assert spec["openapi"] == "3.0.3"
    assert spec["info"]["title"] == "XLENT Scanner API"
    assert "/api/scan-text" in spec["paths"]
    assert "/api/scan-file" in spec["paths"]
    assert "/api/deep-scan" in spec["paths"]
    assert "/microsoft/graph/status" in spec["paths"]
    assert "/microsoft/graph/tags" in spec["paths"]
    assert "/microsoft/graph/resolve-local-file" in spec["paths"]
    assert "/microsoft/graph/tags-for-local-file" in spec["paths"]
    assert "/microsoft/graph/write-scan-metadata" in spec["paths"]
    assert "/microsoft/graph/write-folder-metadata" in spec["paths"]
    assert "ApiKeyAuth" in spec["components"]["securitySchemes"]
    assert docs_response.status_code == 200
    assert "SwaggerUIBundle" in docs_response.get_data(as_text=True)


def test_patch_docx_expands_financial_ai_table_finding_to_cells(tmp_path: Path, monkeypatch):
    from xlent_scanner import redaction_audit

    monkeypatch.setattr(Path, "home", lambda: tmp_path)
    monkeypatch.setattr(redaction_audit, "app_data_dir", lambda: tmp_path)
    downloads = tmp_path / "Downloads"
    downloads.mkdir()

    source = tmp_path / "budget.docx"
    doc = Document()
    table = doc.add_table(rows=4, cols=4)
    rows = [
        ("Item", "Cost (NOK)", "Amount", "Total Cost (NOK)"),
        ("Bread", "30", "2", "60"),
        ("Milk", "20", "1", "20"),
        ("Eggs", "10", "10", "100"),
    ]
    for row, values in zip(table.rows, rows, strict=True):
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
    doc.save(source)

    app_module.app_state.last_path = source
    app_module.app_state.last_result = ScanResult(
        file_name="budget.docx",
        file_size=source.stat().st_size,
        text_length=0,
        text_preview="",
        original_text=(
            "Item | Cost (NOK) | Amount | Total Cost (NOK)\n"
            "Bread | 30 | 2 | 60\n"
            "Milk | 20 | 1 | 20\n"
            "Eggs | 10 | 10 | 100"
        ),
    )

    client = app_module.flask_app.test_client()
    response = client.post(
        "/patch",
        json={
            "indices": [],
            "ai_findings": [
                {
                    "category": "🤖 Budsjettall",
                    "text": "Bread | 30 | 2 | 60",
                    "context": (
                        "Item | Cost (NOK) | Amount | Total Cost (NOK)\n"
                        "Bread | 30 | 2 | 60"
                    ),
                }
            ],
        },
    )

    assert response.status_code == 200
    data = response.get_json()
    assert data["ok"] is True
    assert app_module.app_state.last_anonymized_path == Path(data["path"])
    patched = Document(data["path"])
    bread_cells = [cell.text for cell in patched.tables[0].rows[1].cells]
    milk_cells = [cell.text for cell in patched.tables[0].rows[2].cells]
    assert bread_cells == [
        "Bread",
        "[ANONYMISERT]",
        "[ANONYMISERT]",
        "[ANONYMISERT]",
    ]
    assert milk_cells == ["Milk", "20", "1", "20"]
